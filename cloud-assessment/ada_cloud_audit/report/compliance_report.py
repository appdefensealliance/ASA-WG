"""Generate the ADA Compliance Report .docx from the template.

Template structure (tables):
  Table 0: Header (10 rows x 2 cols) - report number, lab name, app info, summary
  Table 1: Application info (7 rows x 2 cols)
  Table 2: Verdict legend (4 rows x 2 cols) - P, F, NA, INC
  Table 3: Compliance summary (74 rows x 6 cols) - sections by domain with verdict columns
  Table 4: Revision history

Table 3 layout:
  Domain header rows have merged cells in cols 0-1 with domain name, and P/F/NA/INC labels in cols 2-5.
  Section rows have section_id in col 0, section name in col 1, and verdict checkbox cols 2-5 (P, F, NA, INC).
"""

from __future__ import annotations

import re

from docx import Document

from ada_cloud_audit.models import AssessmentReport, Verdict


def _verdict_label(verdict: Verdict) -> str:
    return {
        Verdict.PASS: "P",
        Verdict.FAIL: "F",
        Verdict.NOT_APPLICABLE: "NA",
        Verdict.INCONCLUSIVE: "INC",
    }[verdict]


def _summary_text(report: AssessmentReport) -> str:
    has_fail = any(r.verdict == Verdict.FAIL for r in report.results)
    return "REMEDIATION REQUIRED" if has_fail else "IN COMPLIANCE"


def write_compliance_report(
    report: AssessmentReport,
    template_path: str,
    output_path: str,
) -> None:
    """Generate the Compliance Report .docx from the template."""
    doc = Document(template_path)

    # --- Table 0: Header ---
    header_table = doc.tables[0]
    # Row 1: Lab Name (replace "Lab Name" placeholder)
    _set_cell_text_preserve_format(header_table.rows[1].cells[0], f"\nCloud App and Config Compliance Report\n{report.lab_name}\n")
    # Row 2: Application Name
    header_table.rows[2].cells[1].text = report.app_name
    # Row 3: Application Version
    header_table.rows[3].cells[1].text = report.app_version
    # Row 4: Company
    header_table.rows[4].cells[1].text = report.company
    # Row 7: Summary
    header_table.rows[7].cells[1].text = f"\n{_summary_text(report)}"

    # --- Table 1: Application Info ---
    app_table = doc.tables[1]
    app_table.rows[0].cells[1].text = report.app_name
    app_table.rows[1].cells[1].text = report.app_version
    app_table.rows[3].cells[1].text = report.company

    # --- Table 3: Compliance Summary ---
    summary_table = doc.tables[3]

    # Build section-level verdicts from results
    section_verdicts = {}
    for result in report.results:
        sid = result.section_id
        if sid not in section_verdicts:
            section_verdicts[sid] = []
        section_verdicts[sid].append(result.verdict)

    # Compute each section's aggregate verdict
    computed_verdicts = {}
    for sid, verdicts in section_verdicts.items():
        if any(v == Verdict.FAIL for v in verdicts):
            computed_verdicts[sid] = Verdict.FAIL
        elif any(v == Verdict.INCONCLUSIVE for v in verdicts):
            computed_verdicts[sid] = Verdict.INCONCLUSIVE
        elif all(v in (Verdict.PASS, Verdict.NOT_APPLICABLE) for v in verdicts):
            computed_verdicts[sid] = Verdict.PASS
        else:
            computed_verdicts[sid] = Verdict.NOT_APPLICABLE

    # Walk through the summary table and fill in verdicts
    for row in summary_table.rows:
        cell0_text = row.cells[0].text.strip()
        # Check if this is a section row (has a section ID like "2.8")
        if re.match(r"^\d+\.\d+$", cell0_text):
            section_id = cell0_text
            verdict = computed_verdicts.get(section_id)

            # Clear all verdict columns
            for col in range(2, 6):
                row.cells[col].text = ""

            if verdict:
                # P=col2, F=col3, NA=col4, INC=col5
                verdict_col = {
                    Verdict.PASS: 2,
                    Verdict.FAIL: 3,
                    Verdict.NOT_APPLICABLE: 4,
                    Verdict.INCONCLUSIVE: 5,
                }[verdict]
                row.cells[verdict_col].text = "X"

    doc.save(output_path)


def _set_cell_text_preserve_format(cell, text: str) -> None:
    """Set cell text while trying to preserve existing formatting."""
    if cell.paragraphs:
        cell.paragraphs[0].text = text
    else:
        cell.text = text
