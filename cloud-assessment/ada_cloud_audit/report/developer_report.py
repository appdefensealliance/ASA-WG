"""Generate the ADA Developer Test Report .docx from the template.

Template structure (tables):
  Table 0: Header (10 rows x 2 cols) - same as compliance report
  Table 1: Application info (7 rows x 2 cols)
  Table 2: Verdict legend (4 rows x 2 cols)
  Table 3: Compliance summary (74 rows x 6 cols) - same as compliance report
  Table 4: Revision history
  Tables 5-171: Per-requirement result tables (2 rows x 2 cols each)
    Row 0: ["Result", "<verdict>"]
    Row 1: ["Comment", "<evidence>"]

The spec_id -> table_index mapping is built dynamically by scanning paragraphs for
spec IDs (e.g., "1.2.1") and associating them with the next table after table 4.
"""

from __future__ import annotations

import re

from docx import Document
from docx.oxml.ns import qn

from ada_cloud_audit.models import AssessmentReport, Verdict


def _verdict_text(verdict: Verdict) -> str:
    return {
        Verdict.PASS: "Pass",
        Verdict.FAIL: "Fail",
        Verdict.NOT_APPLICABLE: "NA (Pass)",
        Verdict.INCONCLUSIVE: "Fail",
    }[verdict]


def _summary_text(report: AssessmentReport) -> str:
    has_fail = any(r.verdict == Verdict.FAIL for r in report.results)
    return "REMEDIATION REQUIRED" if has_fail else "IN COMPLIANCE"


def _build_spec_id_to_table_map(doc: Document) -> dict[str, int]:
    """Walk document body elements to map spec_ids to their result table indices."""
    spec_id_pattern = re.compile(r"^(\d+\.\d+\.\d+)\s")
    table_counter = 0
    current_spec_id = None
    mapping = {}

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            # Get full text from all runs
            full_text = ""
            for child in element.iter():
                if child.text:
                    full_text += child.text
                if child.tail:
                    full_text += child.tail
            full_text = full_text.strip()
            match = spec_id_pattern.match(full_text)
            if match:
                current_spec_id = match.group(1)
        elif tag == "tbl":
            if table_counter >= 5 and current_spec_id:
                mapping[current_spec_id] = table_counter
            table_counter += 1

    return mapping


def write_developer_report(
    report: AssessmentReport,
    template_path: str,
    output_path: str,
) -> None:
    """Generate the Developer Test Report .docx from the template."""
    doc = Document(template_path)

    # --- Table 0: Header ---
    header_table = doc.tables[0]
    _set_cell_text(header_table.rows[1].cells[0], f"\nCloud App and Config Developer Test Report\n{report.lab_name}\n")
    header_table.rows[2].cells[1].text = report.app_name
    header_table.rows[3].cells[1].text = report.app_version
    header_table.rows[4].cells[1].text = report.company
    header_table.rows[7].cells[1].text = f"\n{_summary_text(report)}"

    # --- Table 1: Application Info ---
    app_table = doc.tables[1]
    app_table.rows[0].cells[1].text = report.app_name
    app_table.rows[1].cells[1].text = report.app_version
    app_table.rows[3].cells[1].text = report.company

    # --- Table 3: Compliance Summary (same logic as compliance report) ---
    summary_table = doc.tables[3]
    section_verdicts = {}
    for result in report.results:
        sid = result.section_id
        if sid not in section_verdicts:
            section_verdicts[sid] = []
        section_verdicts[sid].append(result.verdict)

    computed_verdicts = {}
    for sid, verdicts in section_verdicts.items():
        if any(v == Verdict.FAIL for v in verdicts):
            computed_verdicts[sid] = Verdict.FAIL
        elif any(v == Verdict.INCONCLUSIVE for v in verdicts):
            computed_verdicts[sid] = Verdict.INCONCLUSIVE
        elif all(v in (Verdict.PASS, Verdict.NOT_APPLICABLE) for v in verdicts):
            computed_verdicts[sid] = Verdict.PASS
        else:
            computed_verdicts[sid] = Verdict.NOT_APPLICABLE

    for row in summary_table.rows:
        cell0_text = row.cells[0].text.strip()
        if re.match(r"^\d+\.\d+$", cell0_text):
            section_id = cell0_text
            verdict = computed_verdicts.get(section_id)
            for col in range(2, 6):
                row.cells[col].text = ""
            if verdict:
                verdict_col = {
                    Verdict.PASS: 2,
                    Verdict.FAIL: 3,
                    Verdict.NOT_APPLICABLE: 4,
                    Verdict.INCONCLUSIVE: 5,
                }[verdict]
                row.cells[verdict_col].text = "X"

    # --- Tables 5+: Per-requirement results ---
    spec_to_table = _build_spec_id_to_table_map(doc)

    # Build a lookup from spec_id to result
    results_by_id = {r.spec_id: r for r in report.results}

    # Fill in each requirement's result table
    for spec_id, table_idx in spec_to_table.items():
        table = doc.tables[table_idx]
        result = results_by_id.get(spec_id)

        if result:
            # Row 0, Col 1: Verdict
            table.rows[0].cells[1].text = _verdict_text(result.verdict)
            # Row 1, Col 1: Evidence/Comment
            table.rows[1].cells[1].text = result.evidence
        else:
            # Non-AWS requirements get NA
            table.rows[0].cells[1].text = "NA (Pass)"
            table.rows[1].cells[1].text = "Not applicable - requirement is for a different cloud provider"

    doc.save(output_path)


def _set_cell_text(cell, text: str) -> None:
    """Set cell text."""
    if cell.paragraphs:
        cell.paragraphs[0].text = text
    else:
        cell.text = text
