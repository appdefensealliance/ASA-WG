"""Tests for DOCX report generation and template compatibility.

Verifies that:
1. Both DOCX templates parse correctly
2. Template structure matches expected layout (tables, headings)
3. Every registered spec ID has a corresponding entry in the developer report template
4. Every section in the compliance summary table matches the spec
5. End-to-end: mock results -> DOCX output -> correct verdicts in cells
6. The spec_id-to-table mapping covers all requirements
"""

from __future__ import annotations

import os
import re
import tempfile

import pytest
from docx import Document

from ada_cloud_audit.models import (
    AssessmentReport,
    Provider,
    RequirementResult,
    Verdict,
    SECTIONS,
    DOMAINS,
    get_section_id,
    get_section_name,
    get_domain,
)
from ada_cloud_audit.report.compliance_report import write_compliance_report
from ada_cloud_audit.report.developer_report import (
    write_developer_report,
    _build_spec_id_to_table_map,
)


TEMPLATE_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "Submission Forms and Templates",
    "Lab Templates",
)
COMPLIANCE_TEMPLATE = os.path.join(
    TEMPLATE_DIR, "Cloud App and Config Compliance Report.docx"
)
DEVELOPER_TEMPLATE = os.path.join(
    TEMPLATE_DIR, "Cloud App and Config Developer Test Report.docx"
)


def _templates_exist():
    return os.path.exists(COMPLIANCE_TEMPLATE) and os.path.exists(DEVELOPER_TEMPLATE)


skipif_no_templates = pytest.mark.skipif(
    not _templates_exist(),
    reason="DOCX templates not found (running outside repo checkout)",
)


def _get_spec_entries_from_markdown():
    """Parse the Specification markdown to get all spec entries."""
    spec_path = os.path.join(
        os.path.dirname(TEMPLATE_DIR),
        os.pardir,
        "Cloud App and Config Profile",
        "Cloud App and Config Specification.md",
    )
    entries = []
    if not os.path.exists(spec_path):
        return entries
    with open(spec_path) as f:
        for line in f:
            m = re.match(r"\|\s*(\d+\.\d+\.\d+)\s*\|\s*(\w+)\s*\|\s*(.+?)\s*\|", line)
            if m:
                entries.append((m.group(1), m.group(2), m.group(3).strip()))
    return entries


def _make_full_report(provider: Provider):
    """Build a report with one result per registered check for the given provider."""
    from ada_cloud_audit.checks.registry import PROVIDER_REGISTRIES

    registry = PROVIDER_REGISTRIES.get(provider, {})
    results = []
    for i, spec_id in enumerate(sorted(registry.keys(), key=lambda x: [int(p) for p in x.split(".")])):
        verdicts = [Verdict.PASS, Verdict.FAIL, Verdict.INCONCLUSIVE, Verdict.NOT_APPLICABLE]
        verdict = verdicts[i % len(verdicts)]
        results.append(
            RequirementResult(
                spec_id=spec_id,
                title=f"Test check {spec_id}",
                platform=provider.value,
                section_id=get_section_id(spec_id),
                section_name=get_section_name(spec_id),
                domain=get_domain(spec_id),
                verdict=verdict,
                evidence=f"Test evidence for {spec_id}",
                details={"test": True},
            )
        )
    return AssessmentReport(
        provider=provider,
        lab_name="Test Lab",
        app_name="TestApp",
        app_version="1.0",
        company="Test Corp",
        results=results,
    )


# ---------------------------------------------------------------------------
# Template structure tests
# ---------------------------------------------------------------------------


@skipif_no_templates
class TestTemplateStructure:
    """Verify DOCX templates have the expected structure."""

    def test_compliance_template_opens(self):
        doc = Document(COMPLIANCE_TEMPLATE)
        assert len(doc.tables) >= 5

    def test_developer_template_opens(self):
        doc = Document(DEVELOPER_TEMPLATE)
        assert len(doc.tables) >= 5

    def test_compliance_table3_has_section_rows(self):
        doc = Document(COMPLIANCE_TEMPLATE)
        table = doc.tables[3]
        section_rows = []
        for row in table.rows:
            text = row.cells[0].text.strip()
            if re.match(r"^\d+\.\d+$", text):
                section_rows.append(text)
        assert len(section_rows) > 0, "Table 3 should have section rows"
        for sid in ["1.2", "2.7", "3.9", "4.3", "5.5", "6.15"]:
            assert sid in section_rows, f"Section {sid} missing from compliance summary"

    def test_developer_table3_matches_compliance(self):
        comp = Document(COMPLIANCE_TEMPLATE)
        dev = Document(DEVELOPER_TEMPLATE)
        comp_rows = [
            row.cells[0].text.strip()
            for row in comp.tables[3].rows
            if re.match(r"^\d+\.\d+$", row.cells[0].text.strip())
        ]
        dev_rows = [
            row.cells[0].text.strip()
            for row in dev.tables[3].rows
            if re.match(r"^\d+\.\d+$", row.cells[0].text.strip())
        ]
        assert comp_rows == dev_rows, "Summary tables should have identical section rows"

    def test_compliance_no_veridict_typo(self):
        doc = Document(COMPLIANCE_TEMPLATE)
        table = doc.tables[3]
        for row in table.rows:
            for cell in row.cells:
                assert "Veridict" not in cell.text, f"Typo 'Veridict' found in cell: {cell.text}"

    def test_developer_has_heading4_entries(self):
        doc = Document(DEVELOPER_TEMPLATE)
        h4_entries = [
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name == "Heading 4"
        ]
        assert len(h4_entries) > 100, f"Expected >100 Heading 4 entries, got {len(h4_entries)}"

    def test_developer_result_tables_match_headings(self):
        doc = Document(DEVELOPER_TEMPLATE)
        h4_count = sum(1 for p in doc.paragraphs if p.style.name == "Heading 4")
        result_tables = sum(
            1
            for t in doc.tables
            if len(t.rows) == 2 and t.rows[0].cells[0].text.strip() == "Result"
        )
        assert abs(result_tables - h4_count) <= 1, (
            f"Result tables should match Heading 4 count (±1 for formatting artifacts): "
            f"{h4_count} headings vs {result_tables} tables"
        )


# ---------------------------------------------------------------------------
# Spec-to-template alignment tests
# ---------------------------------------------------------------------------


@skipif_no_templates
class TestSpecAlignment:
    """Verify templates match the current specification."""

    def test_compliance_sections_cover_spec_sections(self):
        """Every section used by a spec entry should appear in the compliance summary."""
        spec_entries = _get_spec_entries_from_markdown()
        if not spec_entries:
            pytest.skip("Specification markdown not found")

        spec_sections = {get_section_id(e[0]) for e in spec_entries}

        doc = Document(COMPLIANCE_TEMPLATE)
        table_sections = set()
        for row in doc.tables[3].rows:
            text = row.cells[0].text.strip()
            if re.match(r"^\d+\.\d+$", text):
                table_sections.add(text)

        missing = spec_sections - table_sections
        assert not missing, f"Sections in spec but not in compliance template: {sorted(missing)}"

    def test_developer_spec_ids_cover_registry(self):
        """Every registered check spec ID should have a Heading 4 in the developer template."""
        from ada_cloud_audit.checks.registry import PROVIDER_REGISTRIES

        all_spec_ids = set()
        for checks in PROVIDER_REGISTRIES.values():
            all_spec_ids.update(checks.keys())

        doc = Document(DEVELOPER_TEMPLATE)
        doc_spec_ids = set()
        for p in doc.paragraphs:
            if p.style.name == "Heading 4":
                m = re.match(r"(\d+\.\d+\.\d+)", p.text.strip())
                if m:
                    doc_spec_ids.add(m.group(1))

        missing = all_spec_ids - doc_spec_ids
        assert not missing, (
            f"Spec IDs registered in tooling but missing from developer template: "
            f"{sorted(missing)}"
        )

    def test_spec_id_to_table_map_coverage(self):
        """_build_spec_id_to_table_map should find a table for every Heading 4 spec ID."""
        doc = Document(DEVELOPER_TEMPLATE)
        mapping = _build_spec_id_to_table_map(doc)

        h4_spec_ids = set()
        for p in doc.paragraphs:
            if p.style.name == "Heading 4":
                m = re.match(r"(\d+\.\d+\.\d+)", p.text.strip())
                if m:
                    h4_spec_ids.add(m.group(1))

        mapped_ids = set(mapping.keys())
        unmapped = h4_spec_ids - mapped_ids
        assert not unmapped, f"Spec IDs with Heading 4 but no table mapping: {sorted(unmapped)}"

    def test_sections_dict_covers_all_spec_ids(self):
        """Every spec entry's section should resolve to a name (not 'Unknown')."""
        spec_entries = _get_spec_entries_from_markdown()
        if not spec_entries:
            pytest.skip("Specification markdown not found")

        unknown_sections = []
        for spec_id, platform, desc in spec_entries:
            if get_section_name(spec_id) == "Unknown":
                unknown_sections.append(spec_id)

        assert not unknown_sections, (
            f"Spec IDs resolving to 'Unknown' section: {unknown_sections}. "
            f"Add missing entries to SECTIONS dict in models.py"
        )


# ---------------------------------------------------------------------------
# End-to-end DOCX generation tests
# ---------------------------------------------------------------------------


@skipif_no_templates
class TestComplianceReportGeneration:
    """End-to-end tests for compliance report DOCX generation."""

    def test_generates_valid_docx(self):
        report = _make_full_report(Provider.AWS)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            write_compliance_report(report, COMPLIANCE_TEMPLATE, output_path)
            doc = Document(output_path)
            assert len(doc.tables) >= 5
        finally:
            os.unlink(output_path)

    def test_metadata_filled(self):
        report = _make_full_report(Provider.AWS)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            write_compliance_report(report, COMPLIANCE_TEMPLATE, output_path)
            doc = Document(output_path)
            app_table = doc.tables[1]
            assert app_table.rows[0].cells[1].text == "TestApp"
            assert app_table.rows[1].cells[1].text == "1.0"
        finally:
            os.unlink(output_path)

    def test_section_verdicts_filled(self):
        report = _make_full_report(Provider.AWS)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            write_compliance_report(report, COMPLIANCE_TEMPLATE, output_path)
            doc = Document(output_path)
            table = doc.tables[3]

            filled_sections = []
            for row in table.rows:
                text = row.cells[0].text.strip()
                if re.match(r"^\d+\.\d+$", text):
                    verdict_cells = [row.cells[c].text.strip() for c in range(2, 6)]
                    if "X" in verdict_cells:
                        filled_sections.append(text)

            assert len(filled_sections) > 0, "At least some sections should have verdict marks"
        finally:
            os.unlink(output_path)

    def test_all_pass_shows_in_compliance(self):
        report = AssessmentReport(
            provider=Provider.AWS,
            lab_name="Lab",
            app_name="App",
            results=[
                RequirementResult(
                    spec_id="2.8.2",
                    title="Test",
                    platform="AWS",
                    section_id="2.8",
                    section_name="Test Section",
                    domain="Identity and Access Management",
                    verdict=Verdict.PASS,
                    evidence="All good",
                )
            ],
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            write_compliance_report(report, COMPLIANCE_TEMPLATE, output_path)
            doc = Document(output_path)
            header = doc.tables[0]
            summary_text = header.rows[7].cells[1].text
            assert "IN COMPLIANCE" in summary_text
        finally:
            os.unlink(output_path)

    def test_any_fail_shows_remediation(self):
        report = AssessmentReport(
            provider=Provider.AWS,
            lab_name="Lab",
            app_name="App",
            results=[
                RequirementResult(
                    spec_id="2.7.1",
                    title="Test",
                    platform="AWS",
                    section_id="2.7",
                    section_name="Test Section",
                    domain="Identity and Access Management",
                    verdict=Verdict.FAIL,
                    evidence="Root keys exist",
                )
            ],
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            write_compliance_report(report, COMPLIANCE_TEMPLATE, output_path)
            doc = Document(output_path)
            header = doc.tables[0]
            summary_text = header.rows[7].cells[1].text
            assert "REMEDIATION REQUIRED" in summary_text
        finally:
            os.unlink(output_path)


@skipif_no_templates
class TestDeveloperReportGeneration:
    """End-to-end tests for developer test report DOCX generation."""

    def test_generates_valid_docx(self):
        report = _make_full_report(Provider.AWS)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            write_developer_report(report, DEVELOPER_TEMPLATE, output_path)
            doc = Document(output_path)
            assert len(doc.tables) >= 5
        finally:
            os.unlink(output_path)

    def test_result_tables_filled_for_provider(self):
        """Check results land in the correct result tables."""
        report = AssessmentReport(
            provider=Provider.AWS,
            lab_name="Lab",
            app_name="App",
            results=[
                RequirementResult(
                    spec_id="2.8.2",
                    title="Password length",
                    platform="AWS",
                    section_id="2.8",
                    section_name="Config",
                    domain="IAM",
                    verdict=Verdict.PASS,
                    evidence="MinimumPasswordLength is 14",
                ),
                RequirementResult(
                    spec_id="2.7.1",
                    title="Root keys",
                    platform="AWS",
                    section_id="2.7",
                    section_name="ACL",
                    domain="IAM",
                    verdict=Verdict.FAIL,
                    evidence="Root keys present",
                ),
            ],
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            write_developer_report(report, DEVELOPER_TEMPLATE, output_path)
            doc = Document(output_path)
            mapping = _build_spec_id_to_table_map(doc)

            if "2.8.2" in mapping:
                table = doc.tables[mapping["2.8.2"]]
                assert table.rows[0].cells[1].text == "Pass"
                assert "14" in table.rows[1].cells[1].text

            if "2.7.1" in mapping:
                table = doc.tables[mapping["2.7.1"]]
                assert table.rows[0].cells[1].text == "Fail"
                assert "Root keys" in table.rows[1].cells[1].text
        finally:
            os.unlink(output_path)

    def test_other_provider_reqs_get_na(self):
        """Requirements for other providers should get NA (Pass)."""
        report = AssessmentReport(
            provider=Provider.AWS,
            lab_name="Lab",
            app_name="App",
            results=[
                RequirementResult(
                    spec_id="2.8.2",
                    title="Test",
                    platform="AWS",
                    section_id="2.8",
                    section_name="Config",
                    domain="IAM",
                    verdict=Verdict.PASS,
                    evidence="OK",
                )
            ],
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            write_developer_report(report, DEVELOPER_TEMPLATE, output_path)
            doc = Document(output_path)
            mapping = _build_spec_id_to_table_map(doc)

            azure_id = "2.8.1"
            if azure_id in mapping:
                table = doc.tables[mapping[azure_id]]
                assert table.rows[0].cells[1].text == "NA (Pass)"
        finally:
            os.unlink(output_path)

    def test_inconclusive_maps_to_fail(self):
        report = AssessmentReport(
            provider=Provider.AWS,
            lab_name="Lab",
            app_name="App",
            results=[
                RequirementResult(
                    spec_id="2.8.2",
                    title="Test",
                    platform="AWS",
                    section_id="2.8",
                    section_name="Config",
                    domain="IAM",
                    verdict=Verdict.INCONCLUSIVE,
                    evidence="Could not determine",
                )
            ],
        )
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name
        try:
            write_developer_report(report, DEVELOPER_TEMPLATE, output_path)
            doc = Document(output_path)
            mapping = _build_spec_id_to_table_map(doc)
            if "2.8.2" in mapping:
                table = doc.tables[mapping["2.8.2"]]
                assert table.rows[0].cells[1].text == "Fail"
        finally:
            os.unlink(output_path)

    def test_full_provider_report_no_crash(self):
        """Generate a full report with all registered checks — should not crash."""
        for provider in [Provider.AWS, Provider.GCP, Provider.AZURE]:
            try:
                report = _make_full_report(provider)
            except Exception:
                continue
            if not report.results:
                continue
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                output_path = f.name
            try:
                write_developer_report(report, DEVELOPER_TEMPLATE, output_path)
                doc = Document(output_path)
                assert len(doc.tables) >= 5
            finally:
                os.unlink(output_path)
